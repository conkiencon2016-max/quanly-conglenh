
from flask import Flask, render_template, request, redirect, session, send_file
import sqlite3
from datetime import datetime, timedelta
from apscheduler.schedulers.background import BackgroundScheduler
import subprocess
# ===== THÊM NHỮNG IMPORT BỊ THIẾU =====
import shutil
import tempfile
import os
import uuid
import pandas as pd
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.lib import pagesizes
from reportlab.platypus import Image
from flask import jsonify

from flask import send_file
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from io import BytesIO
from flask import after_this_request
from openpyxl.styles import PatternFill
from io import BytesIO


app = Flask(__name__)

app.secret_key = "conglenh_secret_key"

# ⏳ Auto logout 30 phút
app.permanent_session_lifetime = timedelta(minutes=30)

DB = "conglenh.db"

# ================= DATABASE =================
def init_db():
    conn = sqlite3.connect(DB, timeout=30)
    c = conn.cursor()

# ===== BẢNG NGƯỜI ĐI CÔNG TÁC =====
    c.execute("""
    CREATE TABLE IF NOT EXISTS nguoidicongtac(
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    ho_ten TEXT UNIQUE,
    chuc_vu TEXT
)
""")
    # ===== BẢNG NGƯỜI KÝ =====
    c.execute("""
    CREATE TABLE IF NOT EXISTS nguoiky(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ho_ten TEXT
    )
    """)
   # ===== BẢNG CÔNG LỆNH =====
    c.execute("""
    CREATE TABLE IF NOT EXISTS conglenh(
       id INTEGER PRIMARY KEY AUTOINCREMENT,
       nam INTEGER,
       so_thu_tu INTEGER,
       so_cong_lenh TEXT UNIQUE,
       ho_ten TEXT,
       chuc_vu TEXT,
        noi_den TEXT,
       ngay_di TEXT,
       ngay_ve TEXT,
       ngay_ky TEXT,
       nguoi_ky TEXT,
      created_at TEXT,
      UNIQUE(nam, so_thu_tu)
   )
   """)
   
    # Bảng user
    c.execute("""
    CREATE TABLE IF NOT EXISTS users(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE,
        password TEXT,
        role TEXT
    )
    """)

    # Tạo user mặc định
    users = [
        ("admin","123","admin"),
        ("user1","123","user"),
        ("viewer1","123","viewer")
    ]

    for u in users:
        c.execute("SELECT * FROM users WHERE username=?", (u[0],))
        if not c.fetchone():
            c.execute("INSERT INTO users(username,password,role) VALUES(?,?,?)", u)

    conn.commit()
    conn.close()

init_db()

# ================= HELPER =================
def login_required():
    return "user" in session

def check_role(roles):
    return session.get("role") in roles

# ================= LOGIN =================
@app.route("/login", methods=["GET","POST"])
def login():

    if request.method == "POST":

        username = request.form["username"]
        password = request.form["password"]

        conn = sqlite3.connect(DB)
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        c.execute("SELECT * FROM users WHERE username=? AND password=?",
                  (username,password))
        user = c.fetchone()
        conn.close()

        if user:
            session.permanent = True
            session["user"] = user["username"]
            session["role"] = user["role"]
            return redirect("/")
        else:
            return render_template("login.html",
                                   error="Sai tài khoản!")

    return render_template("login.html")

# ================= LOGOUT =================
@app.route("/logout")
def logout():
    session.clear()   # xoá toàn bộ session
    return redirect("/login")

# ================= BẮT BUỘC LOGIN TRƯỚC MỌI TRANG =================
@app.before_request
def require_login():

    allowed_routes = [
        "login",
        "static",
        "api_nguoidicongtac",
        "api_nguoiky"
    ]

    if request.endpoint is None:
        return

    if request.endpoint not in allowed_routes and "user" not in session:
        return redirect("/login")

# ================= HOME =================
@app.route("/")
def home():
    if not login_required():
        return redirect("/login")
    return render_template("home.html")



# ================= NHẬP =================
@app.route("/nhap", methods=["GET", "POST"])
def nhap():

    if "user" not in session:
        return redirect("/login")

    if session.get("role") not in ["admin", "user"]:
        return "Không có quyền!"

    conn = sqlite3.connect(DB, timeout=30)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()

    current_year = datetime.now().year

    # ================= PREVIEW SỐ =================
    def get_next_number():
        row = c.execute("""
            SELECT COALESCE(MAX(so_thu_tu), 0)
            FROM conglenh
            WHERE nam=?
        """, (current_year,)).fetchone()

        return row[0] + 1

    # ================= GET =================
    if request.method == "GET":
        so_display = f"{get_next_number():02}"
        conn.close()
        return render_template(
            "nhapconglenh.html",
            so_cong_lenh=so_display
        )

    action = request.form.get("action")
    so_display = request.form.get("so_cong_lenh")
    so_thu_tu = int(so_display)
    so_full = f"{so_display}/{current_year}"

    # ================= VALIDATE =================
    required = [
        "ho_ten","chuc_vu","noi_den",
        "ngay_di","ngay_ve","ngay_ky","nguoi_ky"
    ]

    for field in required:
        if not request.form.get(field):
            conn.close()
            return render_template(
                "nhapconglenh.html",
                so_cong_lenh=so_display,
                form_data=request.form,
                error="Vui lòng nhập đầy đủ thông tin!"
            )

    if request.form["ngay_ve"] < request.form["ngay_di"]:
        conn.close()
        return render_template(
            "nhapconglenh.html",
            so_cong_lenh=so_display,
            form_data=request.form,
            error="Ngày về không hợp lệ!"
        )

    # ======================================================
    # ======================= SAVE ==========================
    # ======================================================
    if action == "save":
        conn.execute("BEGIN IMMEDIATE")
        existing = c.execute("""
            SELECT id FROM conglenh
            WHERE nam=? AND so_thu_tu=?
        """, (current_year, so_thu_tu)).fetchone()

        if existing:
            # Không báo lỗi — chỉ lấy ID cũ
            saved_id = existing["id"]
        else:
            c.execute("""
                INSERT INTO conglenh(
                    nam, so_thu_tu, so_cong_lenh,
                    ho_ten, chuc_vu, noi_den,
                    ngay_di, ngay_ve, ngay_ky,
                    nguoi_ky, created_at
                ) VALUES (?,?,?,?,?,?,?,?,?,?,?)
            """, (
                current_year,
                so_thu_tu,
                so_full,
                request.form["ho_ten"],
                request.form["chuc_vu"],
                request.form["noi_den"],
                request.form["ngay_di"],
                request.form["ngay_ve"],
                request.form["ngay_ky"],
                request.form["nguoi_ky"],
                datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            ))
            conn.commit()
            saved_id = c.lastrowid

        conn.close()

        return render_template(
            "nhapconglenh.html",
            so_cong_lenh=so_display,
            form_data=request.form,
            saved_id=saved_id,
            success="Đã lưu thành công!"
        )

    # ======================================================
    # ======================= NEXT ==========================
    # ======================================================
    if action == "next":

        conn.execute("BEGIN IMMEDIATE")

        existing = c.execute("""
            SELECT id FROM conglenh
            WHERE nam=? AND so_thu_tu=?
        """, (current_year, so_thu_tu)).fetchone()

        if not existing:
            c.execute("""
                INSERT INTO conglenh(
                    nam, so_thu_tu, so_cong_lenh,
                    ho_ten, chuc_vu, noi_den,
                    ngay_di, ngay_ve, ngay_ky,
                    nguoi_ky, created_at
                ) VALUES (?,?,?,?,?,?,?,?,?,?,?)
            """, (
                current_year,
                so_thu_tu,
                so_full,
                request.form["ho_ten"],
                request.form["chuc_vu"],
                request.form["noi_den"],
                request.form["ngay_di"],
                request.form["ngay_ve"],
                request.form["ngay_ky"],
                request.form["nguoi_ky"],
                datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            ))

        conn.commit()

        # Sinh số mới
        new_number = get_next_number()
        conn.close()

        return render_template(
            "nhapconglenh.html",
            so_cong_lenh=f"{new_number:02}",
            success="Đã chuyển sang công lệnh mới!"
        )

    conn.close()
    return redirect("/nhap")

# ================= DANH SÁCH =================
@app.route("/danhsach")
def danhsach():

    if not login_required():
        return redirect("/login")

    if not check_role(["admin"]):
        return "Không có quyền!"

    so = request.args.get("so", "").strip()
    ten = request.args.get("ten", "").strip()

    conn = sqlite3.connect(DB)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()

    query = "SELECT * FROM conglenh WHERE 1=1"
    params = []

    if so:
        query += " AND LOWER(so_cong_lenh) LIKE LOWER(?)"
        params.append(f"%{so}%")

    if ten:
        query += " AND LOWER(ho_ten) LIKE LOWER(?)"
        params.append(f"%{ten}%")

    query += " ORDER BY id ASC"

    c.execute(query, params)
    data = c.fetchall()
    conn.close()

    return render_template("danhsach.html", data=data)


# ================= ĐỔI MẬT KHẨU =================
@app.route("/change_password", methods=["GET","POST"])
def change_password():

    if not login_required():
        return redirect("/login")

    if request.method == "POST":

        old_pass = request.form["old_password"]
        new_pass = request.form["new_password"]

        conn = sqlite3.connect(DB)
        c = conn.cursor()

        c.execute("SELECT password FROM users WHERE username=?",
                  (session["user"],))
        current = c.fetchone()[0]

        if old_pass != current:
            conn.close()
            return render_template("change_password.html",
                                   error="Mật khẩu cũ không đúng!")

        c.execute("UPDATE users SET password=? WHERE username=?",
                  (new_pass, session["user"]))
        conn.commit()
        conn.close()

        return render_template("change_password.html",
                               success="Đổi mật khẩu thành công!")

    return render_template("change_password.html")
# ================= XÓA =================
@app.route("/xoa/<int:id>")
def xoa(id):
    conn = sqlite3.connect(DB)
    c = conn.cursor()
    c.execute("DELETE FROM conglenh WHERE id=?", (id,))
    conn.commit()
    conn.close()
    return redirect("/danhsach")

# ================= XUẤT EXCEL =================
@app.route("/export_excel")
def export_excel():

    so = request.args.get("so", "").strip()
    ten = request.args.get("ten", "").strip()

    conn = sqlite3.connect("conglenh.db")
    conn.row_factory = sqlite3.Row
    c = conn.cursor()

    query = "SELECT * FROM conglenh WHERE 1=1"
    params = []

    if so:
        query += " AND LOWER(so_cong_lenh) LIKE LOWER(?)"
        params.append(f"%{so}%")

    if ten:
        query += " AND LOWER(ho_ten) LIKE LOWER(?)"
        params.append(f"%{ten}%")

    query += " ORDER BY id DESC"

    c.execute(query, params)
    rows = c.fetchall()
    conn.close()

    # ===== TẠO FILE EXCEL =====
    wb = Workbook()
    ws = wb.active
    ws.title = "Danh sách công lệnh"

    # ===== TIÊU ĐỀ =====
    ws.merge_cells("A1:I1")
    ws["A1"] = "DANH SÁCH CÔNG LỆNH"
    ws["A1"].font = Font(size=16, bold=True)
    ws["A1"].alignment = Alignment(horizontal="center")

    ws.merge_cells("A2:I2")
    ws["A2"] = f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    ws["A2"].alignment = Alignment(horizontal="center")

    headers = [
        "STT",
        "Số công lệnh",
        "Họ tên",
        "Chức vụ",
        "Nơi đến",
        "Ngày đi",
        "Ngày về",
        "Ngày ký",
        "Người ký"
    ]

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col)
        cell.value = header
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center")

    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin")
    )

    for index, row in enumerate(rows, start=5):
        ws.cell(row=index, column=1, value=index - 4)
        ws.cell(row=index, column=2, value=row["so_cong_lenh"])
        ws.cell(row=index, column=3, value=row["ho_ten"])
        ws.cell(row=index, column=4, value=row["chuc_vu"])
        ws.cell(row=index, column=5, value=row["noi_den"])
        ws.cell(row=index, column=6, value=row["ngay_di"])
        ws.cell(row=index, column=7, value=row["ngay_ve"])
        ws.cell(row=index, column=8, value=row["ngay_ky"])
        ws.cell(row=index, column=9, value=row["nguoi_ky"])

        for col in range(1, 10):
            ws.cell(row=index, column=col).border = thin_border

    # ===== AUTO WIDTH =====
    for column_cells in ws.iter_cols(min_row=4):
        length = max(len(str(cell.value)) if cell.value else 0 for cell in column_cells)
        ws.column_dimensions[column_cells[0].column_letter].width = length + 3

    output = BytesIO()
    wb.save(output)
    output.seek(0)

    return send_file(
        output,
        as_attachment=True,
        download_name="Danh_sach_cong_lenh.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

# ================= XUẤT EXCEL THỐNG KÊ =================

@app.route("/export_excel_thongke", methods=["POST"])
def export_excel_thongke():

    tu = request.form.get("tu_ngay","")
    den = request.form.get("den_ngay","")
    ten = request.form.get("ho_ten","")

    conn = sqlite3.connect(DB)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()

    query = "SELECT * FROM conglenh WHERE 1=1"
    params = []

    if tu and den:
        query += " AND ngay_di BETWEEN ? AND ?"
        params.extend([tu, den])

    if ten:
        query += " AND ho_ten LIKE ?"
        params.append(f"%{ten}%")

    c.execute(query, params)
    rows = c.fetchall()
    conn.close()

    # ===== TẠO FILE EXCEL =====
    wb = Workbook()
    ws = wb.active
    ws.title = "Bao Cao Cong Lenh"

    # ===== TIÊU ĐỀ =====
    ws.merge_cells("A1:F1")
    ws["A1"] = "BÁO CÁO THỐNG KÊ CÔNG LỆNH"
    ws["A1"].font = Font(size=16, bold=True)
    ws["A1"].alignment = Alignment(horizontal="center")

    # ===== HEADER =====
    headers = ["STT","Số CL","Họ tên","Nơi đến","Ngày đi","Ngày về"]
    ws.append([])
    ws.append(headers)

    header_row = ws.max_row

    for col in range(1, len(headers)+1):
        cell = ws.cell(row=header_row, column=col)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="1F4E78",
                                end_color="1F4E78",
                                fill_type="solid")
        cell.alignment = Alignment(horizontal="center")

    # ===== DỮ LIỆU =====
    thin = Side(border_style="thin")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for index, row in enumerate(rows, start=1):

        ngay_di = datetime.strptime(row["ngay_di"], "%Y-%m-%d").strftime("%d/%m/%Y")
        ngay_ve = datetime.strptime(row["ngay_ve"], "%Y-%m-%d").strftime("%d/%m/%Y")

        ws.append([
            index,
            row["so_cong_lenh"],
            row["ho_ten"],
            row["noi_den"],
            ngay_di,
            ngay_ve
        ])

        for col in range(1,7):
            ws.cell(row=ws.max_row, column=col).border = border

    # ===== TỔNG =====
    ws.append([])
    ws.append([f"Tổng số công lệnh: {len(rows)}"])
    ws.cell(row=ws.max_row, column=1).font = Font(bold=True)

    # ===== AUTO WIDTH =====
    for col in ws.columns:
        max_length = 0
        column = col[0].column
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[get_column_letter(column)].width = adjusted_width

    # ===== LƯU FILE =====
    file_path = os.path.join(
        tempfile.gettempdir(),
        f"thongke_{uuid.uuid4().hex}.xlsx"
    )

    wb.save(file_path)

    return send_file(file_path, as_attachment=True)
# ================= XUẤT EXCEL THỐNG KÊ HÀM MỚI =================


@app.route("/export_thongke", methods=["GET"])
def export_thongke():

    tu_ngay = request.args.get("tu_ngay", "")
    den_ngay = request.args.get("den_ngay", "")
    ten = request.args.get("ten", "")

    conn = sqlite3.connect("conglenh.db")
    conn.row_factory = sqlite3.Row
    c = conn.cursor()

    query = "SELECT * FROM conglenh WHERE 1=1"
    params = []

    if tu_ngay:
        query += " AND date(ngay_di) >= date(?)"
        params.append(tu_ngay)

    if den_ngay:
        query += " AND date(ngay_di) <= date(?)"
        params.append(den_ngay)

    if ten:
        query += " AND LOWER(ho_ten) LIKE LOWER(?)"
        params.append(f"%{ten}%")

    query += " ORDER BY ngay_di DESC"

    c.execute(query, params)
    rows = c.fetchall()
    conn.close()

    wb = Workbook()
    ws = wb.active
    ws.title = "Thống kê công lệnh"

    ws.merge_cells("A1:I1")
    ws["A1"] = "THỐNG KÊ CÔNG LỆNH"
    ws["A1"].font = Font(size=16, bold=True)
    ws["A1"].alignment = Alignment(horizontal="center")

    ws.merge_cells("A2:I2")
    ws["A2"] = f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    ws["A2"].alignment = Alignment(horizontal="center")

    headers = [
        "STT",
        "Số công lệnh",
        "Họ tên",
        "Chức vụ",
        "Nơi đến",
        "Ngày đi",
        "Ngày về",
        "Ngày ký",
        "Người ký"
    ]

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col)
        cell.value = header
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center")

    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin")
    )

    for index, row in enumerate(rows, start=5):
        ws.cell(row=index, column=1, value=index - 4)
        ws.cell(row=index, column=2, value=row["so_cong_lenh"])
        ws.cell(row=index, column=3, value=row["ho_ten"])
        ws.cell(row=index, column=4, value=row["chuc_vu"])
        ws.cell(row=index, column=5, value=row["noi_den"])
        ws.cell(row=index, column=6, value=row["ngay_di"])
        ws.cell(row=index, column=7, value=row["ngay_ve"])
        ws.cell(row=index, column=8, value=row["ngay_ky"])
        ws.cell(row=index, column=9, value=row["nguoi_ky"])

        for col in range(1, 10):
            ws.cell(row=index, column=col).border = thin_border

    for column_cells in ws.iter_cols(min_row=4):
        length = max(len(str(cell.value)) if cell.value else 0 for cell in column_cells)
        ws.column_dimensions[column_cells[0].column_letter].width = length + 3

    output = BytesIO()
    wb.save(output)
    output.seek(0)

    return send_file(
        output,
        as_attachment=True,
        download_name="Thong_ke_cong_lenh.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


# ================= THỐNG KÊ =================
@app.route("/thongke", methods=["GET","POST"])
def thongke():

    data = []

    if request.method=="POST":

        tu = request.form["tu_ngay"]
        den = request.form["den_ngay"]
        ten = request.form["ho_ten"]

        conn = sqlite3.connect(DB)
        conn.row_factory = sqlite3.Row
        c = conn.cursor()

        query = "SELECT * FROM conglenh WHERE 1=1"
        params = []

        if tu and den:
            query += " AND ngay_di BETWEEN ? AND ?"
            params.extend([tu, den])

        if ten:
            query += " AND ho_ten LIKE ?"
            params.append(f"%{ten}%")

        c.execute(query, params)
        data = c.fetchall()
        conn.close()

    return render_template("thongke.html", data=data)
# ================= SỬA =================
@app.route("/sua/<int:id>", methods=["GET","POST"])
def sua(id):

    conn = sqlite3.connect(DB)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()

    if request.method == "POST":

        fields = ["ho_ten","chuc_vu","noi_den",
                  "ngay_di","ngay_ve","ngay_ky","nguoi_ky"]

        for f in fields:
            if not request.form.get(f):
                conn.close()
                return "Thiếu dữ liệu!"

        if request.form["ngay_ve"] < request.form["ngay_di"]:
            conn.close()
            return "Ngày về không hợp lệ!"

        c.execute("""
            UPDATE conglenh SET
                ho_ten=?,
                chuc_vu=?,
                noi_den=?,
                ngay_di=?,
                ngay_ve=?,
                ngay_ky=?,
                nguoi_ky=?
            WHERE id=?
        """,(
            request.form["ho_ten"],
            request.form["chuc_vu"],
            request.form["noi_den"],
            request.form["ngay_di"],
            request.form["ngay_ve"],
            request.form["ngay_ky"],
            request.form["nguoi_ky"],
            id
        ))

        conn.commit()
        conn.close()
        return redirect("/danhsach")

    c.execute("SELECT * FROM conglenh WHERE id=?", (id,))
    data = c.fetchone()
    conn.close()

    return render_template("sua.html", row=data)
# ================= HÀM REPLACE CHUẨN =================
def replace_advanced(doc, data):

    bold_fields = ["{{HOTEN}}", "{{NGUOIKY}}"]

    # ===== XỬ LÝ PARAGRAPH =====
    for p in doc.paragraphs:

        original_text = p.text
        new_text = original_text

        for key, value in data.items():
            new_text = new_text.replace(key, value)

        if original_text != new_text:

            p.clear()

            current_pos = 0

            for key in bold_fields:
                if key in original_text:

                    value = data.get(key, "")
                    parts = new_text.split(value)

                    for i, part in enumerate(parts):

                        if part:
                            p.add_run(part)

                        if i < len(parts) - 1:
                            run = p.add_run(value)
                            run.bold = True

                    break
            else:
                p.add_run(new_text)

    # ===== XỬ LÝ TABLE =====
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:

                    original_text = p.text
                    new_text = original_text

                    for key, value in data.items():
                        new_text = new_text.replace(key, value)

                    if original_text != new_text:

                        p.clear()

                        for key in bold_fields:
                            if key in original_text:

                                value = data.get(key, "")
                                parts = new_text.split(value)

                                for i, part in enumerate(parts):

                                    if part:
                                        p.add_run(part)

                                    if i < len(parts) - 1:
                                        run = p.add_run(value)
                                        run.bold = True

                                break
                        else:
                            p.add_run(new_text) 

# ================= IN word =================

@app.route("/in_word_form", methods=["POST"])
def in_word_form():

    try:
        doc = Document("CL2026.docx")

        # ================= LẤY SỐ CÔNG LỆNH =================
        so_display = request.form.get("so_cong_lenh")

        if not so_display:
            return "Thiếu số công lệnh!"

        # ================= XỬ LÝ NGÀY =================
        try:
            d1 = datetime.strptime(request.form["ngay_di"], "%Y-%m-%d")
            d2 = datetime.strptime(request.form["ngay_ve"], "%Y-%m-%d")
            d3 = datetime.strptime(request.form["ngay_ky"], "%Y-%m-%d")
        except ValueError:
            return "Lỗi định dạng ngày!"

        data = {
            "{{HOTEN}}": request.form.get("ho_ten", ""),
            "{{CHUCVU}}": request.form.get("chuc_vu", ""),
            "{{NOIDEN}}": request.form.get("noi_den", ""),
            "{{SOCONG}}": f"{int(so_display):02}",
            "{{NGUOIKY}}": request.form.get("nguoi_ky", ""),
            "{{NOIDEN_BANG}}": request.form.get("noi_den", ""),
            "{{NGAYDI}}": f"{d1.day:02}",
            "{{THANGDI}}": f"{d1.month:02}",
            "{{NAMDI}}": str(d1.year),
            "{{NGAYVE}}": f"{d2.day:02}",
            "{{THANGVE}}": f"{d2.month:02}",
            "{{NAMVE}}": str(d2.year),
            "{{NGAYKY}}": f"{d3.day:02}",
            "{{THANGKY}}": f"{d3.month:02}",
            "{{NAMKY}}": str(d3.year),
        }

        replace_advanced(doc, data)

        temp_path = os.path.join(
            tempfile.gettempdir(),
            f"conglenh_{uuid.uuid4().hex}.docx"
        )

        doc.save(temp_path)

        return send_file(
            temp_path,
            as_attachment=True,
            download_name="conglenh.docx"
        )

    except Exception:
        return "Không thể tạo file Word"

# ================= IN WORD TỪ DANH SÁCH =================
@app.route("/in_word/<int:id>")
def in_word(id):

    conn = sqlite3.connect(DB)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()

    row = c.execute(
        "SELECT * FROM conglenh WHERE id=?",
        (id,)
    ).fetchone()

    conn.close()

    if not row:
        return "Không tìm thấy công lệnh!"

    # ===== XỬ LÝ NGÀY =====
    d1 = datetime.strptime(row["ngay_di"], "%Y-%m-%d")
    d2 = datetime.strptime(row["ngay_ve"], "%Y-%m-%d")
    d3 = datetime.strptime(row["ngay_ky"], "%Y-%m-%d")

    data = {
        "{{HOTEN}}": row["ho_ten"],
        "{{CHUCVU}}": row["chuc_vu"],
        "{{NOIDEN}}": row["noi_den"],
        "{{SOCONG}}": f"{row['so_thu_tu']:02}",
        "{{NGUOIKY}}": row["nguoi_ky"],
        "{{NOIDEN_BANG}}": row["noi_den"],
        "{{NGAYDI}}": f"{d1.day:02}",
        "{{THANGDI}}": f"{d1.month:02}",
        "{{NAMDI}}": str(d1.year),
        "{{NGAYVE}}": f"{d2.day:02}",
        "{{THANGVE}}": f"{d2.month:02}",
        "{{NAMVE}}": str(d2.year),
        "{{NGAYKY}}": f"{d3.day:02}",
        "{{THANGKY}}": f"{d3.month:02}",
        "{{NAMKY}}": str(d3.year),
    }

    doc = Document("CL2026.docx")

    replace_advanced(doc, data)

    temp_path = os.path.join(
        tempfile.gettempdir(),
        f"conglenh_{uuid.uuid4().hex}.docx"
    )

    doc.save(temp_path)

    return send_file(
        temp_path,
        as_attachment=True,
        download_name="conglenh.docx"
    )

# ===== API NGƯỜI ĐI CÔNG TÁC =====
@app.route("/api/nguoidicongtac")
def api_nguoidicongtac():
    conn = sqlite3.connect(DB)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT * FROM nguoidicongtac ORDER BY ho_ten ASC")
    data = [dict(row) for row in c.fetchall()]
    conn.close()
    return jsonify(data)

# ===== API NGƯỜI KÝ =====
@app.route("/api/nguoiky")
def api_nguoiky():
    conn = sqlite3.connect(DB)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT * FROM nguoiky ORDER BY ho_ten ASC")
    data = [dict(row) for row in c.fetchall()]
    conn.close()
    return jsonify(data)
# ================= QUẢN LÝ NGƯỜI ĐI CÔNG TÁC =================
@app.route("/quanly_nguoidicongtac")
def quanly_nguoidicongtac():
    if not login_required():
        return redirect("/login")

    if not check_role(["admin"]):
        return "Không có quyền!"

    conn = sqlite3.connect(DB)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()

    c.execute("SELECT * FROM nguoidicongtac ORDER BY ho_ten ASC")
    data = c.fetchall()

    conn.close()

    return render_template(
        "quanly_nguoidicongtac.html",
        data=data
    )

# ================= THÊM NGƯỜI ĐI CÔNG TÁC =================


@app.route("/them_nguoidicongtac", methods=["POST"])
def them_nguoidicongtac():

    if not check_role(["admin"]):
        return "Không có quyền!"

    ho_ten = request.form["ho_ten"]
    chuc_vu = request.form["chuc_vu"]

    conn = sqlite3.connect(DB)
    c = conn.cursor()

    c.execute("""
        INSERT INTO nguoidicongtac(ho_ten, chuc_vu)
        VALUES (?,?)
    """, (ho_ten, chuc_vu))

    conn.commit()
    conn.close()

    return redirect("/quanly_nguoidicongtac")

# ================= SỬA NGƯỜI ĐI CÔNG TÁC =================


@app.route("/sua_nguoidicongtac/<int:id>", methods=["GET","POST"])
def sua_nguoidicongtac(id):

    if not login_required():
        return redirect("/login")

    if not check_role(["admin"]):
        return "Không có quyền!"

    conn = sqlite3.connect(DB)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()

    if request.method == "POST":
        ho_ten = request.form["ho_ten"].strip()
        chuc_vu = request.form["chuc_vu"].strip()

        # ===== KIỂM TRA TRÙNG TÊN (TRỪ CHÍNH NÓ) =====
        c.execute("""
            SELECT id FROM nguoidicongtac
            WHERE ho_ten=? AND id!=?
        """, (ho_ten, id))

        if c.fetchone():
            conn.close()
            return "Tên cán bộ đã tồn tại!"

        # ===== UPDATE =====
        c.execute("""
            UPDATE nguoidicongtac
            SET ho_ten=?, chuc_vu=?
            WHERE id=?
        """, (ho_ten, chuc_vu, id))

        conn.commit()
        conn.close()

        return redirect("/quanly_nguoidicongtac")

    # ===== GET =====
    c.execute("SELECT * FROM nguoidicongtac WHERE id=?", (id,))
    row = c.fetchone()
    conn.close()

    if not row:
        return "Không tìm thấy cán bộ!"

    return render_template("sua_nguoidicongtac.html", row=row)

# ================= XÓA NGƯỜI ĐI CÔNG TÁC =================
@app.route("/xoa_nguoidicongtac/<int:id>")
def xoa_nguoidicongtac(id):

    if not check_role(["admin"]):
        return "Không có quyền!"

    conn = sqlite3.connect(DB)
    c = conn.cursor()

    c.execute("DELETE FROM nguoidicongtac WHERE id=?", (id,))
    conn.commit()
    conn.close()

    return redirect("/quanly_nguoidicongtac")
# ================= backup =================
# ================= BACKUP DATABASE =================
def backup_job():
    try:
        subprocess.run(["python", "backup_drive.py"], check=True)
        print("Backup Google Drive OK")
    except Exception as e:
        print("Backup Google Drive lỗi:", e)

# ================= AUTO BACKUP DATABASE =================
def auto_backup():

    try:

        os.makedirs("backups", exist_ok=True)

        today = datetime.now().strftime("%Y%m%d")

        backup_file = f"backups/conglenh_{today}.db"

        shutil.copy(DB, backup_file)

        print("Backup OK:", backup_file)

        # giữ 30 file gần nhất
        files = sorted(os.listdir("backups"))

        if len(files) > 30:

            for f in files[:-30]:
                os.remove(os.path.join("backups", f))

    except Exception as e:

        print("Backup error:", e)
# ================= download_backup =================
@app.route("/download_backup/<filename>")
def download_backup(filename):

    if session.get("role") != "admin":
        return "Không có quyền!"

    path = os.path.join("backups", filename)

    if not os.path.exists(path):
        return "File không tồn tại!"

    return send_file(path, as_attachment=True)
# ================= restore_backup =================
@app.route("/restore_backup/<filename>")
def restore_backup(filename):

    if session.get("role") != "admin":
        return "Không có quyền!"

    backup_file = os.path.join("backups", filename)

    if not os.path.exists(backup_file):
        return "File backup không tồn tại!"

    shutil.copy(backup_file, DB)

    return "Khôi phục database thành công! Hãy reload trang."
# ================= backup_manager =================

@app.route("/backup_manager")
def backup_manager():

    os.makedirs("backups", exist_ok=True)

    files = sorted(os.listdir("backups"), reverse=True)

    return render_template(
        "backup_manager.html",
        files=files
    )
# ================= backup_now =================

@app.route("/backup_now")
def backup_now():

    if session.get("role") != "admin":
        return "Không có quyền!"

    auto_backup()

    return redirect("/backup_manager")
# ================= SCHEDULER =================

scheduler = BackgroundScheduler()
scheduler.add_job(auto_backup, 'cron', hour=2)
# chạy mỗi ngày lúc 02:00 sáng
scheduler.add_job(backup_job, 'cron', hour=2, minute=0)

if os.environ.get("RUN_MAIN") != "true":
    scheduler.start()

if __name__ == "__main__":

    port = int(os.environ.get("PORT", 10000))

    print("Server đang chạy...")
    print("Backup database mỗi ngày lúc 02:00")

    app.run(host="0.0.0.0", port=port)






